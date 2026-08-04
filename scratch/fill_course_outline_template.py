from __future__ import annotations

import hashlib
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from build_template_matched_documents import ANALYSIS_ROWS, PROJECT_ROWS


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = Path("/private/tmp/nos-course-outline-template/converted/เค้าโครงการสอน-e1.docx")
OUTPUT = ROOT / "output" / "docx" / "เค้าโครงการสอน_31901-2002_ภาคเรียน1_2569.docx"
INTERMEDIATE = Path("/private/tmp/nos-course-outline-template/filled-python-docx.docx")

REFERENCE_SHA256 = "8428ef4b3f2f64093639d2c54364efcf560ad0a4c67552ca2a3d0e5e51fa6998"

FONT = "Tahoma"
COURSE_CODE = "31901-2002"
COURSE_NAME = "ระบบปฏิบัติการเครื่องแม่ข่าย"
COURSE_ENGLISH = "Network Operating System for Server"
CREDITS = "1-4-3"
COLLEGE = "วิทยาลัยอาชีวศึกษานครปฐม"
TEACHER = "นายสุรชาติ ลิ้มรัตนพันธ์"
LEVEL = "ปวส.1"
SEMESTER = "1/2569"
WEEKLY_HOURS = 5
TOTAL_HOURS = 75
COURSEWORK_SCORE = 35


OBJECTIVE_LINES = [
    "1. เข้าใจโมเดล OSI 7 Layer, TCP/IP และหลักการทำงานของเครือข่ายคอมพิวเตอร์พื้นฐาน",
    "2. เข้าใจการติดตั้งและใช้งานระบบปฏิบัติการเครื่องแม่ข่าย รวมถึงแพ็กเก็ตที่สนับสนุนบริการเครือข่าย",
    "3. มีทักษะติดตั้งและบริหารจัดการ Proxmox Hypervisor และเครื่องเสมือน",
    "4. มีทักษะติดตั้งบริการ DNS, Web, Database, DHCP, File Sharing, Proxy, AAA และ Container Platform",
    "5. สามารถ Deploy Web Application บนเครื่องแม่ข่ายด้วยแนวคิด DevOps เบื้องต้น",
    "6. มีเจตคติและกิจนิสัยที่ดีในการปฏิบัติงานอย่างละเอียดรอบคอบและรับผิดชอบ",
]

COMPETENCY_LINES = [
    "1. อธิบายหลักการทำงานของโมเดล OSI 7 Layer และ TCP/IP Protocol ได้",
    "2. ติดตั้งระบบปฏิบัติการเครื่องแม่ข่ายและแพ็กเก็ตบริการเครือข่ายตามขั้นตอนได้",
    "3. ติดตั้งและบริหารจัดการ Proxmox VE Hypervisor พร้อมสร้าง VM หรือ LXC ได้",
    "4. ติดตั้งและทดสอบบริการ DNS, Web, Database, DHCP, File Sharing, Proxy, AAA และ Container Platform ได้",
    "5. Deploy Web Application ด้วย Git, Node.js, MariaDB และ Nginx ตามแนวคิด DevOps เบื้องต้นได้",
    "6. ใช้งาน ตรวจสอบ และแก้ปัญหาระบบปฏิบัติการเครื่องแม่ข่ายได้ตรงตามวัตถุประสงค์",
]

DESCRIPTION_LINES = [
    "ศึกษาและปฏิบัติเกี่ยวกับโมเดล OSI 7 Layer และ TCP/IP Protocol ระบบปฏิบัติการบนเครื่องแม่ข่าย",
    "การเข้าหัวสาย LAN ตามมาตรฐาน TIA/EIA การติดตั้ง Proxmox Hypervisor และการสร้างหรือจัดการเครื่องเสมือน",
    "การตั้งค่าพื้นฐาน การจัดการผู้ใช้งาน Firewall และ SSH รวมถึงการตรวจสอบระบบเครือข่ายและพอร์ตบริการ",
    "การติดตั้ง DNS Server, Web Server, Database Server, DHCP Server และ File and Resource Sharing Server",
    "การประยุกต์ Proxy Server, AAA Server, Container Platform และ IoT Platform ตามความต้องการของระบบ",
    "การใช้ SSL Certificate เพื่อปกป้องการสื่อสาร และการตั้งค่า Nginx Reverse Proxy สำหรับ Web Application",
    "การ Deploy Web Application ด้วย Git, Node.js และ MariaDB ตามแนวคิด DevOps เบื้องต้น",
]


ANALYSIS_COMPACT = [
    {
        "unit": item["unit"],
        "name": item["name"],
        "competency": item["competency"],
        "objective": item["objective"],
        "content": item["content"],
        "activity": item["activity"],
        "media": item["media"],
        "assessment": item["assessment"],
        "hours": item["hours"],
    }
    for item in ANALYSIS_ROWS
]
for analysis_item in ANALYSIS_COMPACT:
    if analysis_item["unit"] == "12":
        analysis_item["assessment"] = "Project 15 คะแนน; Rubric; Demo; Runbook; การทำงานเป็นทีม"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def set_run_font(run, size: float, bold: bool = False) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)


def scale_visible_template_fonts(doc: Document) -> None:
    size_map = {20.0: 15.0, 18.0: 14.0, 17.0: 13.0, 16.0: 11.5, 14.0: 10.0, 11.0: 8.0, 3.0: 3.0}
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            current = run.font.size.pt if run.font.size else 14.0
            target = size_map.get(round(current, 1), 10.0 if current <= 14 else 11.5)
            set_run_font(run, target, bool(run.bold))


def write_paragraph(
    paragraph,
    text: str,
    size: float = 10.0,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    left_indent_mm: float | None = None,
    line_spacing: float = 1.0,
) -> None:
    paragraph.clear()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = line_spacing
    if left_indent_mm is not None:
        paragraph.paragraph_format.left_indent = Mm(left_indent_mm)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold)


def set_cell_margins(cell, top: int = 35, start: int = 45, bottom: int = 35, end: int = 45) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def write_cell(
    cell,
    text: str,
    size: float,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    line_spacing: float = 0.92,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = line_spacing
    run = paragraph.add_run(str(text))
    set_run_font(run, size, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def write_schedule_cell(cell, title: str, line1: str, line2: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 0.9
    run = paragraph.add_run(title)
    set_run_font(run, 8.1, True)
    for line in (line1, line2):
        run = paragraph.add_run("\n" + line)
        set_run_font(run, 7.6, False)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=25, start=45, bottom=25, end=45)


def remove_excess_schedule_rows(table) -> None:
    # The converted source contains 21 teaching-row placeholders. Keep 15 plus the total row.
    while len(table.rows) > 17:
        table._tbl.remove(table.rows[-2]._tr)


def fill_schedule(table) -> None:
    remove_excess_schedule_rows(table)
    set_repeat_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    headers = ["สัปดาห์\nที่", "หน่วย\nที่", "ชื่อหน่วย/รายการสอน", "จำนวน\nชั่วโมง", "คะแนนเก็บ"]
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, 9.3, True, WD_ALIGN_PARAGRAPH.CENTER, 0.95)

    for row_index, item in enumerate(PROJECT_ROWS, start=1):
        week, unit, title, line1, line2, hours, score = item
        row = table.rows[row_index]
        set_row_cant_split(row)
        write_cell(row.cells[0], str(week), 8.6, True, WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(row.cells[1], str(unit), 8.6, False, WD_ALIGN_PARAGRAPH.CENTER)
        write_schedule_cell(row.cells[2], title, line1, line2)
        write_cell(row.cells[3], str(hours), 8.6, False, WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(row.cells[4], "" if score is None else str(score), 8.6, False, WD_ALIGN_PARAGRAPH.CENTER)

    total_row = table.rows[-1]
    set_row_cant_split(total_row)
    write_cell(total_row.cells[0], "รวมจำนวนชั่วโมง/ภาคเรียน", 9.2, True, WD_ALIGN_PARAGRAPH.RIGHT)
    write_cell(total_row.cells[3], str(TOTAL_HOURS), 9.2, True, WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(total_row.cells[4], str(COURSEWORK_SCORE), 9.2, True, WD_ALIGN_PARAGRAPH.CENTER)


def merge_vertical_group(table, start: int, end: int) -> list:
    cells = []
    for column in range(9):
        merged = table.cell(start, column).merge(table.cell(end, column))
        cells.append(merged)
    set_row_cant_split(table.rows[start])
    return cells


def fill_analysis(table) -> None:
    for header_row in (0, 22):
        set_repeat_header(table.rows[header_row])
        set_row_cant_split(table.rows[header_row])
        for cell in table.rows[header_row].cells:
            # Preserve the source labels while making Thai reliably render.
            existing = cell.text
            write_cell(cell, existing, 7.6, True, WD_ALIGN_PARAGRAPH.CENTER, 0.9)

    groups = [
        (1, 3),
        (4, 6),
        (7, 9),
        (10, 12),
        (13, 15),
        (16, 18),
        (19, 21),
        (23, 26),
        (27, 30),
        (31, 34),
        (35, 38),
        (39, 42),
        (43, 47),
    ]

    for item, (start, end) in zip(ANALYSIS_COMPACT, groups):
        cells = merge_vertical_group(table, start, end)
        values = [
            item["unit"],
            item["name"],
            item["competency"],
            item["objective"],
            item["content"],
            item["activity"],
            item["media"],
            item["assessment"],
            item["hours"],
        ]
        for column, (cell, value) in enumerate(zip(cells, values)):
            align = WD_ALIGN_PARAGRAPH.CENTER if column in (0, 3, 8) else WD_ALIGN_PARAGRAPH.LEFT
            write_cell(cell, str(value), 6.7, column in (0, 1), align, 0.86)


def compress_empty_signature_spacers(doc: Document) -> None:
    for index in list(range(55, 59)) + list(range(63, 68)) + list(range(72, 78)):
        paragraph = doc.paragraphs[index]
        paragraph.clear()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 0.7
        run = paragraph.add_run("")
        set_run_font(run, 2.0, False)


def fill_document(doc: Document) -> None:
    scale_visible_template_fonts(doc)

    p = doc.paragraphs
    write_paragraph(p[7], f"โครงการสอน ภาคเรียนที่ {SEMESTER}", 12.0, True, WD_ALIGN_PARAGRAPH.CENTER)
    write_paragraph(
        p[8],
        f"ชื่อวิชา {COURSE_NAME} ({COURSE_ENGLISH})    รหัสวิชา {COURSE_CODE}    (ท-ป-น) {CREDITS}",
        10.2,
        False,
        WD_ALIGN_PARAGRAPH.LEFT,
    )
    write_paragraph(
        p[9],
        f"เวลาเรียน {WEEKLY_HOURS} ชั่วโมง/สัปดาห์    รวม {TOTAL_HOURS} ชั่วโมง/ภาคเรียน",
        10.2,
        False,
        WD_ALIGN_PARAGRAPH.LEFT,
    )
    write_paragraph(
        p[10],
        f"ระดับชั้น {LEVEL}    ครูผู้สอน {TEACHER}    สถานศึกษา {COLLEGE}",
        9.8,
        False,
        WD_ALIGN_PARAGRAPH.LEFT,
    )

    objective_blocks = [OBJECTIVE_LINES[0], OBJECTIVE_LINES[1], OBJECTIVE_LINES[2], "\n".join(OBJECTIVE_LINES[3:])]
    competency_blocks = [COMPETENCY_LINES[0], COMPETENCY_LINES[1], COMPETENCY_LINES[2], "\n".join(COMPETENCY_LINES[3:])]
    for index, text in zip(range(13, 17), objective_blocks):
        write_paragraph(p[index], text, 9.2, False, WD_ALIGN_PARAGRAPH.LEFT, 5.0, 0.92)
    for index, text in zip(range(18, 22), competency_blocks):
        write_paragraph(p[index], text, 9.2, False, WD_ALIGN_PARAGRAPH.LEFT, 5.0, 0.92)
    for index, text in zip(range(23, 30), DESCRIPTION_LINES):
        write_paragraph(p[index], text, 9.0, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 5.0 if index == 23 else 0.0, 0.92)

    write_paragraph(p[37], f"ชื่อวิชา {COURSE_NAME}    รหัสวิชา {COURSE_CODE}    (ท-ป-น) {CREDITS}", 10.0)
    write_paragraph(p[38], f"เวลาเรียน {WEEKLY_HOURS} ชั่วโมง/สัปดาห์    รวม {TOTAL_HOURS} ชั่วโมง/ภาคเรียน", 10.0)
    fill_schedule(doc.tables[0])
    write_paragraph(
        p[40],
        "สัดส่วนคะแนน: คะแนนเก็บ/ใบงาน 35 คะแนน | Project 15 คะแนน | สอบปฏิบัติ 10 คะแนน | สอบปลายภาค 20 คะแนน | จิตพิสัย 20 คะแนน",
        8.7,
        True,
        WD_ALIGN_PARAGRAPH.CENTER,
    )

    write_paragraph(p[46], f"วิชา {COURSE_NAME}    รหัสวิชา {COURSE_CODE}    จำนวน (ท-ป-น) {CREDITS}", 11.0, True)
    write_paragraph(p[47], f"เวลาเรียน {WEEKLY_HOURS} ชั่วโมง : สัปดาห์    รวม {TOTAL_HOURS} ชั่วโมง : ภาคเรียน", 11.0, True)
    fill_analysis(doc.tables[1])

    write_paragraph(p[52], "ลงชื่อ.............................................................ครูผู้สอน", 10.0, False, WD_ALIGN_PARAGRAPH.CENTER)
    write_paragraph(p[53], f"({TEACHER})", 10.0, False, WD_ALIGN_PARAGRAPH.CENTER)
    write_paragraph(p[54], "ครูประจำ/ครูอัตราจ้าง สาขาวิชาเทคโนโลยีสารสนเทศ", 10.0, False, WD_ALIGN_PARAGRAPH.CENTER)
    compress_empty_signature_spacers(doc)

    # Use the existing page break before the approval area only if the table does not already force one.
    p[52].paragraph_format.page_break_before = True


def preserve_reference_package(modified: Path, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(modified, "r") as edited_zip:
        document_xml = edited_zip.read("word/document.xml")
    with zipfile.ZipFile(REFERENCE, "r") as source_zip, zipfile.ZipFile(output, "w") as out_zip:
        for info in source_zip.infolist():
            data = document_xml if info.filename == "word/document.xml" else source_zip.read(info.filename)
            out_zip.writestr(info, data)


def main() -> None:
    if sha256(REFERENCE) != REFERENCE_SHA256:
        raise RuntimeError("Converted template hash changed; distill the template again before editing.")

    shutil.copy2(REFERENCE, INTERMEDIATE)
    doc = Document(INTERMEDIATE)
    fill_document(doc)
    doc.save(INTERMEDIATE)
    preserve_reference_package(INTERMEDIATE, OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
