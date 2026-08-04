from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Image, Paragraph, Table, TableStyle

import build_official_teaching_documents as source


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "output" / "docx"
OUT_PDF = ROOT / "output" / "pdf"
LOGO = ROOT / "output" / "college-logo-npvc.png"

FONT = "Tahoma"
PDF_FONT = "Template-Tahoma"
PDF_BOLD = "Template-Tahoma-Bold"
FONT_REGULAR_PATH = "/System/Library/Fonts/Supplemental/Tahoma.ttf"
FONT_BOLD_PATH = "/System/Library/Fonts/Supplemental/Tahoma Bold.ttf"


PROJECT_ROWS = [
    (1, 1, "พื้นฐานระบบปฏิบัติการเครื่องแม่ข่าย", "Server/Client, NOS, OSI 7 Layer และ TCP/IP", "เครือข่าย LAN และมาตรฐานสาย TIA/EIA-568A/B", 5, 2),
    (2, 2, "ติดตั้ง Ubuntu Server และคำสั่งพื้นฐาน", "ติดตั้งระบบและจัดการ Package, File และ Directory", "กำหนด IP Address เวลา และตรวจสอบระบบ", 5, 3),
    (3, 3, "บริการเครือข่าย DNS และ DHCP", "หลักการ IP Address, DHCP DORA และ Lease", "DNS Record และการตรวจสอบชื่อเครื่อง", 5, 2),
    (4, 4, "Server Hardware และ Proxmox VE", "Server, Virtualization, Storage และ Network Bridge", "ติดตั้ง Proxmox VE และใช้งาน Web UI", 5, 3),
    (5, 5, "Proxmox Container และ Ubuntu Server Stack", "สร้าง LXC Container และกำหนดเครือข่าย", "Git, Nginx, MariaDB, Node.js และ Reverse Proxy", 5, 4),
    (6, 6, "User, Permission และ SSH", "User/Group, chmod, chown และ sudo", "SSH, Key Authentication และการตรวจสอบ Log", 5, 3),
    (7, 7, "Server Network และ UFW Firewall", "ตรวจ IP, Route, Port และ Socket", "UFW, IP Whitelisting, Nmap และ Security Log", 5, 3),
    (8, 7, "ทบทวนระบบ Server Stack", "ตรวจความพร้อม Container และบริการที่ติดตั้ง", "วิเคราะห์ปัญหา Service, Port และ Permission", 5, None),
    (9, 8, "FTP Server ด้วย vsftpd", "Control/Data Connection, Port 21 และ Passive Port", "ติดตั้ง vsftpd สร้าง User/Folder และกำหนดสิทธิ์", 5, 3),
    (10, 8, "ทดสอบ FTP ด้วย FileZilla", "เปิด Firewall สำหรับ Port 21 และ Passive Port Range", "Upload/Download และวิเคราะห์ Message Log", 5, 3),
    (11, 9, "HTTPS, SSL/TLS และ Nginx Reverse Proxy", "Certificate, HTTPS Port 443 และ Reverse Proxy", "เข้ารหัสการเชื่อมต่อและตรวจ Configuration", 5, 3),
    (12, 10, "Deploy Web Application และ DevOps เบื้องต้น", "Git, Node.js, MariaDB, Nginx และ Environment", "นำ Application ขึ้นใช้งาน ทดสอบ Browser และจัดทำ Runbook", 5, 3),
    (13, 11, "File Sharing, Proxy, AAA และ IoT", "Samba, Proxy/ACL, RADIUS และ MQTT Broker", "ทดสอบบริการจาก Client และตรวจสอบสิทธิ์", 5, 3),
    (14, 12, "Project บูรณาการระบบเครื่องแม่ข่าย", "ออกแบบและติดตั้งบริการบน Proxmox Container", "ทดสอบระบบ จัดทำ Runbook และนำเสนอ", 5, None),
    (15, 13, "ประเมินผลปลายภาคและสรุปการเรียนรู้", "สอบปฏิบัติการติดตั้งและแก้ปัญหา Server", "สอบปลายภาคและสรุปสมรรถนะรายวิชา", 5, None),
]

ANALYSIS_ROWS = [dict(item) for item in source.ANALYSIS]
ANALYSIS_ROWS[9] = {
    "unit": "10",
    "name": "Deploy Web Application และ DevOps เบื้องต้น",
    "competency": "ติดตั้งและนำ Web Application ขึ้นใช้งาน พร้อมจัดการบริการพื้นฐานอย่างเป็นระบบ",
    "objective": "ข้อ 4, 5 และ 6",
    "content": "Git; Node.js; MariaDB; Nginx Reverse Proxy; Environment; Service/Process",
    "activity": "Clone/Update Code; ติดตั้ง Dependency; ตั้งค่า Environment; เริ่ม Service; ทดสอบผ่าน Browser",
    "media": "Git; Node.js; MariaDB; Nginx; Browser; คู่มือ Lab",
    "assessment": "ตรวจ Web Application; ตรวจ Service/Port; ใบงาน Deploy",
    "hours": 5,
}


pdfmetrics.registerFont(TTFont(PDF_FONT, FONT_REGULAR_PATH))
pdfmetrics.registerFont(TTFont(PDF_BOLD, FONT_BOLD_PATH))
pdfmetrics.registerFontFamily(
    PDF_FONT,
    normal=PDF_FONT,
    bold=PDF_BOLD,
    italic=PDF_FONT,
    boldItalic=PDF_BOLD,
)


def set_run_font(run, size=10, bold=False, underline=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rpr.rFonts.set(qn(key), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline


def set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = line


def set_cell_margins(cell, top=35, start=45, bottom=35, end=45):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_mm):
    cell.width = Mm(width_mm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(round(width_mm * 56.692913)))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths_mm, border=True):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    total = round(sum(widths_mm) * 56.692913)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_mm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 56.692913)))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_mm[index])
            set_cell_margins(cell)
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single" if border else "none")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")
    if not border:
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_borders = tc_pr.find(qn("w:tcBorders"))
                if tc_borders is None:
                    tc_borders = OxmlElement("w:tcBorders")
                    tc_pr.append(tc_borders)
                for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    node = tc_borders.find(qn(f"w:{edge}"))
                    if node is None:
                        node = OxmlElement(f"w:{edge}")
                        tc_borders.append(node)
                    node.set(qn("w:val"), "none")


def exact_row_height(row, height_mm):
    row.height = Mm(height_mm)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def write_cell(cell, text, size=8, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, line=0.95):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph(p, align, line)
    run = p.add_run(str(text))
    set_run_font(run, size, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def write_topic_cell(cell, title, line1, line2, size=7.2):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT, 0.88)
    run = p.add_run(title)
    set_run_font(run, size, True)
    for text in (line1, line2):
        run = p.add_run("\n" + text)
        set_run_font(run, size, False)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_metadata_row(doc, widths, cells, font_size=10):
    table = doc.add_table(rows=1, cols=len(widths))
    configure_table(table, widths, border=False)
    for cell, parts in zip(table.rows[0].cells, cells):
        cell.text = ""
        p = cell.paragraphs[0]
        set_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT, 1.0)
        for text, underline in parts:
            run = p.add_run(text)
            set_run_font(run, font_size, False, underline)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    exact_row_height(table.rows[0], 7)
    return table


def configure_doc_styles(doc, size=10):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(size)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        normal._element.get_or_add_rPr().rFonts.set(qn(key), FONT)


def build_project_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    configure_doc_styles(doc)
    doc.core_properties.title = f"โครงการสอน {source.COURSE_CODE} {source.COURSE_NAME}"
    doc.core_properties.author = source.TEACHER
    doc.core_properties.subject = source.SEMESTER

    p = doc.add_paragraph()
    set_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, 1.0)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run("โครงการสอน/แผนการจัดการเรียนรู้ตลอดภาคเรียน/ กิจกรรมการเรียนการสอน")
    set_run_font(run, 13.5, True)

    add_metadata_row(
        doc,
        [86, 48, 36],
        [
            [("ชื่อวิชา ", False), (source.COURSE_NAME, True)],
            [("รหัสวิชา ", False), (source.COURSE_CODE, True)],
            [("(ท-ป-น) ", False), ("1-4-3", True)],
        ],
        9.3,
    )
    add_metadata_row(
        doc,
        [86, 84],
        [
            [("เวลาเรียน ", False), ("5", True), (" ชั่วโมง/สัปดาห์", False)],
            [("รวม ", False), ("75", True), (" ชั่วโมง/ภาคเรียน", False)],
        ],
        9.3,
    )

    table = doc.add_table(rows=1, cols=5)
    widths = [16, 14.5, 94, 20, 25.5]
    configure_table(table, widths, border=True)
    headers = ["สัปดาห์\nที่", "หน่วย\nที่", "ชื่อหน่วย/รายการสอน", "จำนวน\nชั่วโมง", "คะแนนเก็บ"]
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, 9.2, True, WD_ALIGN_PARAGRAPH.CENTER, 0.95)
    exact_row_height(table.rows[0], 17)

    for week, unit, title, line1, line2, hours, score in PROJECT_ROWS:
        row = table.add_row()
        configure_table(table, widths, border=True)
        exact_row_height(row, 12.9)
        write_cell(row.cells[0], week, 8.2, False, WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(row.cells[1], unit, 8.2, False, WD_ALIGN_PARAGRAPH.CENTER)
        write_topic_cell(row.cells[2], title, line1, line2, 7.0)
        write_cell(row.cells[3], hours, 8.2, False, WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(row.cells[4], "" if score is None else score, 8.2, False, WD_ALIGN_PARAGRAPH.CENTER)

    row = table.add_row()
    configure_table(table, widths, border=True)
    exact_row_height(row, 9)
    merged = row.cells[0].merge(row.cells[2])
    write_cell(merged, "รวมจำนวนชั่วโมง/ภาคเรียน", 9, True, WD_ALIGN_PARAGRAPH.LEFT)
    write_cell(row.cells[3], "75", 9, True, WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(row.cells[4], "35", 9, True, WD_ALIGN_PARAGRAPH.CENTER)

    path = OUT_DOCX / f"โครงการสอน_{source.COURSE_CODE}_ภาคเรียน1_2569.docx"
    doc.save(path)
    return path


def build_analysis_docx():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(8)
    section.bottom_margin = Mm(8)
    section.left_margin = Mm(6)
    section.right_margin = Mm(6)
    configure_doc_styles(doc, 7)
    doc.core_properties.title = f"ตารางวิเคราะห์การจัดการเรียนรู้ {source.COURSE_CODE}"
    doc.core_properties.author = source.TEACHER
    doc.core_properties.subject = source.SEMESTER

    p = doc.add_paragraph()
    set_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, 1.0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(LOGO), width=Mm(23.5))
    p = doc.add_paragraph()
    set_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER, 1.0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("ตารางวิเคราะห์การจัดการเรียนรู้")
    set_run_font(run, 12.5, True)

    add_metadata_row(
        doc,
        [105, 90, 90],
        [
            [("วิชา ", False), (source.COURSE_NAME, True)],
            [("รหัสวิชา ", False), (source.COURSE_CODE, True)],
            [("จำนวน (ท-ป-น) ", False), ("1-4-3", True)],
        ],
        8.5,
    )
    add_metadata_row(
        doc,
        [142, 143],
        [
            [("เวลาเรียน ", False), ("5", True), (" ชั่วโมง : สัปดาห์", False)],
            [("รวม ", False), ("75", True), (" ชั่วโมง : ภาคเรียน", False)],
        ],
        8.5,
    )

    widths = [15.7, 37.8, 40.7, 40.7, 43.5, 40.7, 21.9, 30.4, 13.6]
    table = doc.add_table(rows=1, cols=9)
    configure_table(table, widths, border=True)
    headers = [
        "หน่วย\nที่",
        "ชื่อหน่วย",
        "สมรรถนะวิชาชีพ",
        "จุดประสงค์รายวิชา",
        "เนื้อหาสาระ (โดยย่อ)",
        "กิจกรรมการเรียนรู้",
        "สื่อ\nการเรียนรู้",
        "วิธีการประเมิน",
        "เวลา\n(ชม.)",
    ]
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, 7.2, True, WD_ALIGN_PARAGRAPH.CENTER, 0.9)
    exact_row_height(table.rows[0], 13.5)

    for item in ANALYSIS_ROWS:
        row = table.add_row()
        configure_table(table, widths, border=True)
        exact_row_height(row, 8.2)
        values = [
            item["unit"],
            item["name"],
            item["competency"],
            item["objective"],
            item["content"],
            item["activity"],
            item["media"],
            item["assessment"],
            item["hours"],
        ]
        for index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if index in (0, 3, 8) else WD_ALIGN_PARAGRAPH.LEFT
            write_cell(row.cells[index], value, 5.2, index in (0, 1), align, 0.78)

    path = OUT_DOCX / f"ตารางวิเคราะห์การจัดการเรียนรู้_{source.COURSE_CODE}_ภาคเรียน1_2569.docx"
    doc.save(path)
    return path


PROJECT_CELL = ParagraphStyle(
    "ProjectCell",
    fontName=PDF_FONT,
    fontSize=6.6,
    leading=7.6,
    alignment=TA_LEFT,
    wordWrap="CJK",
)
PROJECT_CENTER = ParagraphStyle(
    "ProjectCenter",
    parent=PROJECT_CELL,
    fontSize=8.0,
    leading=9.0,
    alignment=TA_CENTER,
)
ANALYSIS_CELL = ParagraphStyle(
    "AnalysisCell",
    fontName=PDF_FONT,
    fontSize=4.8,
    leading=5.5,
    alignment=TA_LEFT,
    wordWrap="CJK",
)
ANALYSIS_CENTER = ParagraphStyle(
    "AnalysisCenter",
    parent=ANALYSIS_CELL,
    alignment=TA_CENTER,
)


def ptext(text, style, bold=False):
    value = escape(str(text)).replace("\n", "<br/>")
    if bold:
        value = f"<b>{value}</b>"
    return Paragraph(value, style)


def draw_underlined_field(pdf, x, y, label, value, width, size=9):
    pdf.setFont(PDF_FONT, size)
    pdf.drawString(x, y, label)
    offset = pdfmetrics.stringWidth(label, PDF_FONT, size) + 2
    pdf.setFont(PDF_FONT, size)
    pdf.drawString(x + offset, y, value)
    pdf.setLineWidth(0.35)
    pdf.line(x + offset, y - 1.3, x + width, y - 1.3)


def build_project_pdf():
    path = OUT_PDF / f"โครงการสอน_{source.COURSE_CODE}_ภาคเรียน1_2569.pdf"
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    pdf.setTitle(f"โครงการสอน {source.COURSE_CODE} {source.COURSE_NAME}")
    pdf.setAuthor(source.TEACHER)
    pdf.setFont(PDF_BOLD, 13.5)
    pdf.drawCentredString(width / 2, height - 22 * mm, "โครงการสอน/แผนการจัดการเรียนรู้ตลอดภาคเรียน/ กิจกรรมการเรียนการสอน")

    y1 = height - 34 * mm
    draw_underlined_field(pdf, 20 * mm, y1, "ชื่อวิชา ", source.COURSE_NAME, 83 * mm, 9.2)
    draw_underlined_field(pdf, 105 * mm, y1, "รหัสวิชา ", source.COURSE_CODE, 42 * mm, 9.2)
    draw_underlined_field(pdf, 150 * mm, y1, "(ท-ป-น) ", "1-4-3", 40 * mm, 9.2)
    y2 = height - 44 * mm
    draw_underlined_field(pdf, 20 * mm, y2, "เวลาเรียน ", "5 ชั่วโมง/สัปดาห์", 83 * mm, 9.2)
    draw_underlined_field(pdf, 105 * mm, y2, "รวม ", "75 ชั่วโมง/ภาคเรียน", 85 * mm, 9.2)

    headers = ["สัปดาห์\nที่", "หน่วย\nที่", "ชื่อหน่วย/รายการสอน", "จำนวน\nชั่วโมง", "คะแนนเก็บ"]
    data = [[ptext(value, PROJECT_CENTER, True) for value in headers]]
    for week, unit, title, line1, line2, hours, score in PROJECT_ROWS:
        topic = Paragraph(
            f"<b>{escape(title)}</b><br/>{escape(line1)}<br/>{escape(line2)}",
            PROJECT_CELL,
        )
        data.append(
            [
                ptext(week, PROJECT_CENTER),
                ptext(unit, PROJECT_CENTER),
                topic,
                ptext(hours, PROJECT_CENTER),
                ptext("" if score is None else score, PROJECT_CENTER),
            ]
        )
    data.append(
        [
            ptext("รวมจำนวนชั่วโมง/ภาคเรียน", PROJECT_CELL, True),
            "",
            "",
            ptext("75", PROJECT_CENTER, True),
            ptext("35", PROJECT_CENTER, True),
        ]
    )
    row_heights = [17 * mm] + [12.9 * mm] * 15 + [9 * mm]
    table = Table(
        data,
        colWidths=[16 * mm, 14.5 * mm, 94 * mm, 20 * mm, 25.5 * mm],
        rowHeights=row_heights,
    )
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("SPAN", (0, -1), (2, -1)),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 1.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1.5),
            ]
        )
    )
    table.wrapOn(pdf, 170 * mm, sum(row_heights))
    table.drawOn(pdf, 20 * mm, 29 * mm)
    pdf.showPage()
    pdf.save()
    return path


def build_analysis_pdf():
    page_size = landscape(A4)
    path = OUT_PDF / f"ตารางวิเคราะห์การจัดการเรียนรู้_{source.COURSE_CODE}_ภาคเรียน1_2569.pdf"
    pdf = canvas.Canvas(str(path), pagesize=page_size)
    width, height = page_size
    pdf.setTitle(f"ตารางวิเคราะห์การจัดการเรียนรู้ {source.COURSE_CODE}")
    pdf.setAuthor(source.TEACHER)

    logo = Image(str(LOGO), width=23.5 * mm, height=23.5 * mm)
    logo.wrapOn(pdf, 23.5 * mm, 23.5 * mm)
    logo.drawOn(pdf, width / 2 - 11.75 * mm, height - 37.5 * mm)
    pdf.setFont(PDF_BOLD, 12)
    pdf.drawCentredString(width / 2, height - 43 * mm, "ตารางวิเคราะห์การจัดการเรียนรู้")
    y1 = height - 50 * mm
    draw_underlined_field(pdf, 42 * mm, y1, "วิชา ", source.COURSE_NAME, 100 * mm, 8.5)
    draw_underlined_field(pdf, 145 * mm, y1, "รหัสวิชา ", source.COURSE_CODE, 70 * mm, 8.5)
    draw_underlined_field(pdf, 218 * mm, y1, "จำนวน (ท-ป-น) ", "1-4-3", 60 * mm, 8.5)
    y2 = height - 59 * mm
    draw_underlined_field(pdf, 70 * mm, y2, "เวลาเรียน ", "5 ชั่วโมง : สัปดาห์", 95 * mm, 8.5)
    draw_underlined_field(pdf, 170 * mm, y2, "รวม ", "75 ชั่วโมง : ภาคเรียน", 78 * mm, 8.5)

    headers = [
        "หน่วย\nที่",
        "ชื่อหน่วย",
        "สมรรถนะวิชาชีพ",
        "จุดประสงค์รายวิชา",
        "เนื้อหาสาระ (โดยย่อ)",
        "กิจกรรมการเรียนรู้",
        "สื่อ\nการเรียนรู้",
        "วิธีการประเมิน",
        "เวลา\n(ชม.)",
    ]
    data = [[ptext(value, ANALYSIS_CENTER, True) for value in headers]]
    for item in ANALYSIS_ROWS:
        values = [
            item["unit"],
            item["name"],
            item["competency"],
            item["objective"],
            item["content"],
            item["activity"],
            item["media"],
            item["assessment"],
            item["hours"],
        ]
        data.append(
            [
                ptext(value, ANALYSIS_CENTER if index in (0, 3, 8) else ANALYSIS_CELL, index in (0, 1))
                for index, value in enumerate(values)
            ]
        )
    widths = [15.7, 37.8, 40.7, 40.7, 43.5, 40.7, 21.9, 30.4, 13.6]
    row_heights = [13.5 * mm] + [8.2 * mm] * len(ANALYSIS_ROWS)
    table = Table(data, colWidths=[value * mm for value in widths], rowHeights=row_heights)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 1.5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 1.5),
                ("TOPPADDING", (0, 0), (-1, -1), 1.0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1.0),
            ]
        )
    )
    table.wrapOn(pdf, 285 * mm, sum(row_heights))
    table.drawOn(pdf, 6 * mm, 24 * mm)
    pdf.showPage()
    pdf.save()
    return path


if __name__ == "__main__":
    OUT_DOCX.mkdir(parents=True, exist_ok=True)
    OUT_PDF.mkdir(parents=True, exist_ok=True)
    for output in (build_project_docx(), build_analysis_docx(), build_project_pdf(), build_analysis_pdf()):
        print(output)
