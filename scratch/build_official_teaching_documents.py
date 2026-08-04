from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path("/Users/surachartlimrattanaphun/Desktop/NOS/NOS")
OUT_DOCX = ROOT / "output" / "docx"
LOGO = ROOT / "output" / "college-logo-npvc.png"

COURSE_CODE = "31901-2002"
COURSE_NAME = "ระบบปฏิบัติการเครื่องแม่ข่าย"
COURSE_NAME_EN = "Network Operating System for Server"
COURSE_CREDITS = "1-4-3"
TEACHER = "นายสุรชาติ ลิ้มรัตนพันธ์"
LEVEL = "ปวส.1"
SEMESTER = "ภาคเรียนที่ 1 ปีการศึกษา 2569"
COLLEGE = "วิทยาลัยอาชีวศึกษานครปฐม"
WEEKLY_HOURS = 5
TOTAL_HOURS = 75
FONT = "Arial Unicode MS"

COURSE_DESCRIPTION = (
    "ศึกษาและปฏิบัติเกี่ยวกับโมเดล OSI 7 Layer และ TCP/IP Protocol ระบบปฏิบัติการบนเครื่องแม่ข่าย "
    "การเข้าหัวสาย LAN การติดตั้ง Proxmox Hypervisor การสร้างและจัดการเครื่องเสมือน การตั้งค่าพื้นฐาน "
    "การจัดการผู้ใช้งาน Firewall, SSH, DNS Server, Web Server, Database Server, DHCP Server, "
    "File and Resource Sharing Server, Proxy Server, AAA Server, Container Platform, IoT Platform, "
    "SSL Certificate และการ Deploy Web Application ด้วยแนวคิด DevOps เบื้องต้น"
)

SCHEDULE = [
    {
        "week": "1",
        "unit": "1",
        "title": "พื้นฐานระบบปฏิบัติการเครื่องแม่ข่ายและเครือข่าย",
        "topics": [
            "บทบาทของ Server และ Client",
            "OSI 7 Layer, TCP/IP และโปรโตคอลสำคัญ",
            "มาตรฐานสาย LAN TIA/EIA-568A/B และการทดสอบสาย",
        ],
        "hours": 5,
        "score": 2,
    },
    {
        "week": "2",
        "unit": "2",
        "title": "ติดตั้ง Ubuntu Server และใช้งานคำสั่งพื้นฐาน",
        "topics": [
            "ติดตั้งระบบปฏิบัติการและตั้งค่าเครือข่ายหลังติดตั้ง",
            "จัดการไฟล์ โฟลเดอร์ และแพ็กเกจด้วยคำสั่ง Linux",
            "ตรวจสอบ IP Address และสถานะระบบ",
        ],
        "hours": 5,
        "score": 3,
    },
    {
        "week": "3",
        "unit": "3",
        "title": "บริการเครือข่ายพื้นฐาน DNS และ DHCP",
        "topics": [
            "หลักการแจก IP Address และกระบวนการ DHCP DORA",
            "หลักการแปลงชื่อของ DNS และชนิดของ Record",
            "ทบทวนคำสั่ง Linux CLI ที่จำเป็นต่อการดูแล Server",
        ],
        "hours": 5,
        "score": 2,
    },
    {
        "week": "4",
        "unit": "4",
        "title": "สถาปัตยกรรม Server และ Proxmox VE",
        "topics": [
            "ส่วนประกอบ Hardware ของ Server และ Virtualization",
            "ติดตั้ง Proxmox VE และเข้าใช้งาน Web UI",
            "ตั้งค่า Network Bridge และ Storage เบื้องต้น",
        ],
        "hours": 5,
        "score": 3,
    },
    {
        "week": "5",
        "unit": "5",
        "title": "Proxmox Container และ Ubuntu Server Stack",
        "topics": [
            "สร้างและจัดการ LXC Container",
            "ติดตั้ง Git, Nginx, MariaDB และ Node.js",
            "ตั้งค่า Nginx Reverse Proxy และตรวจสอบบริการ",
        ],
        "hours": 5,
        "score": 4,
    },
    {
        "week": "6",
        "unit": "6",
        "title": "การจัดการผู้ใช้ สิทธิ์ และ SSH",
        "topics": [
            "สร้าง User และ Group พร้อมกำหนดสิทธิ์ chmod และ chown",
            "เชื่อมต่อและรับส่งไฟล์ด้วย SSH",
            "ใช้ SSH Key และปรับความปลอดภัยของบริการ SSH",
        ],
        "hours": 5,
        "score": 3,
    },
    {
        "week": "7",
        "unit": "7",
        "title": "เครือข่าย Server การตรวจสอบพอร์ต และ UFW Firewall",
        "topics": [
            "ตรวจสอบ IP, Route, Socket และ Port ที่เปิดใช้งาน",
            "กำหนด UFW Rule, Default Policy และ IP Whitelisting",
            "ตรวจสอบความปลอดภัยด้วย Nmap และวิเคราะห์ Log",
        ],
        "hours": 5,
        "score": 3,
    },
    {
        "week": "8",
        "unit": "7",
        "title": "ทบทวนและบูรณาการระบบ Server Stack",
        "topics": [
            "ทบทวน Nginx, MariaDB, Node.js, SSH และ Firewall",
            "ตรวจความพร้อมของ Container ก่อนติดตั้งบริการ FTP",
            "ฝึกวิเคราะห์ปัญหาจากสถานะ Service, Port และ Permission",
        ],
        "hours": 5,
        "score": 0,
    },
    {
        "week": "9",
        "unit": "8",
        "title": "FTP Server และการติดตั้ง vsftpd",
        "topics": [
            "หลักการ FTP, Control Connection และ Data Connection",
            "ติดตั้งและตั้งค่า vsftpd บน Proxmox Container",
            "สร้างผู้ใช้ โฟลเดอร์ และกำหนด Passive Port",
        ],
        "hours": 5,
        "score": 3,
    },
    {
        "week": "10",
        "unit": "8",
        "title": "ทดสอบ FTP ด้วย FileZilla และแก้ปัญหา",
        "topics": [
            "เปิด Firewall สำหรับ Port 21 และ Passive Port Range",
            "ทดสอบ Upload และ Download ด้วย FileZilla",
            "วิเคราะห์ Message Log และแก้ปัญหา Permission หรือ Directory Listing",
        ],
        "hours": 5,
        "score": 3,
    },
    {
        "week": "11",
        "unit": "9",
        "title": "Nginx Reverse Proxy และการเข้ารหัส SSL/TLS",
        "topics": [
            "หลักการ Reverse Proxy, HTTPS และ Certificate",
            "ออก Self-signed Certificate และตั้งค่า Port 443",
            "ทดสอบการเชื่อมต่อแบบเข้ารหัสและตรวจสอบ Configuration",
        ],
        "hours": 5,
        "score": 3,
    },
    {
        "week": "12",
        "unit": "10",
        "title": "Docker และการ Deploy Web Application",
        "topics": [
            "เปรียบเทียบ Container กับ Virtual Machine",
            "ติดตั้ง Docker และจัดการ Image, Container และ Volume",
            "Deploy Web Application ด้วย Docker Compose",
        ],
        "hours": 5,
        "score": 3,
    },
    {
        "week": "13",
        "unit": "11",
        "title": "File Sharing, Proxy, AAA และ IoT Platform",
        "topics": [
            "สร้างพื้นที่แชร์ไฟล์ด้วย Samba และกำหนดสิทธิ์",
            "ศึกษาหลักการ Proxy Server และ AAA Server",
            "สาธิต MQTT Broker และการรับส่งข้อความ IoT",
        ],
        "hours": 5,
        "score": 3,
    },
    {
        "week": "14",
        "unit": "12",
        "title": "Project บูรณาการระบบเครื่องแม่ข่าย",
        "topics": [
            "ออกแบบระบบบน Proxmox Container ตามโจทย์",
            "ติดตั้งบริการ Web, Database, Firewall และบริการเสริม",
            "จัดทำแผนผังระบบ เอกสาร Runbook และนำเสนอผลการทำงาน",
        ],
        "hours": 5,
        "score": 15,
    },
    {
        "week": "15",
        "unit": "13",
        "title": "ประเมินผลปลายภาคและสรุปผลการเรียนรู้",
        "topics": [
            "สอบปลายภาค 20 คะแนน",
            "สอบปฏิบัติติดตั้งและแก้ปัญหา Server 10 คะแนน",
            "ประเมินจิตพิสัยตลอดภาคเรียน 20 คะแนน และสรุป Reflection",
        ],
        "hours": 5,
        "score": 50,
    },
]

ANALYSIS = [
    {
        "unit": "1",
        "name": "พื้นฐาน NOS, OSI, TCP/IP และระบบเครือข่าย",
        "competency": "อธิบายโครงสร้างเครือข่ายและปฏิบัติงานสายสัญญาณตามมาตรฐาน",
        "objective": "ข้อ 1, 2 และ 6",
        "content": "บทบาท Server/Client; OSI 7 Layer; TCP/IP; Protocol; TIA/EIA-568A/B; LAN Tester",
        "activity": "อภิปรายจากสถานการณ์; จับคู่ Protocol กับ Layer; เข้าหัวและทดสอบสาย LAN",
        "media": "สไลด์; สาย UTP; RJ-45; คีม; LAN Tester",
        "assessment": "ใบงาน; แบบทดสอบ; ตรวจชิ้นงานสาย LAN",
        "hours": 5,
    },
    {
        "unit": "2",
        "name": "Ubuntu Server และ Linux CLI",
        "competency": "ติดตั้งระบบปฏิบัติการเครื่องแม่ข่ายและตั้งค่าพื้นฐานตามขั้นตอน",
        "objective": "ข้อ 2 และ 6",
        "content": "ติดตั้ง Ubuntu Server; Network; Package; File/Directory; System Command",
        "activity": "สาธิตการติดตั้ง; ปฏิบัติบนเครื่องหรือ VM; ทำภารกิจคำสั่งเป็นลำดับ",
        "media": "Ubuntu ISO; Server/VM; คู่มือคำสั่ง; ใบงาน",
        "assessment": "Checklist การติดตั้ง; ใบงาน CLI; สังเกตการปฏิบัติ",
        "hours": 5,
    },
    {
        "unit": "3",
        "name": "บริการ DNS และ DHCP",
        "competency": "อธิบายและตรวจสอบบริการกำหนดค่าที่อยู่และชื่อในเครือข่าย",
        "objective": "ข้อ 1, 2 และ 4",
        "content": "DHCP DORA; Lease; Scope; DNS Hierarchy; Record; คำสั่งตรวจสอบเครือข่าย",
        "activity": "วิเคราะห์ลำดับ DORA; จำลองการแปลงชื่อ; ทดลองคำสั่งบน Server",
        "media": "Animation; แผนผังเครือข่าย; Ubuntu Container; ใบงาน",
        "assessment": "แบบทดสอบ; ใบงานวิเคราะห์; ตรวจผลคำสั่ง",
        "hours": 5,
    },
    {
        "unit": "4",
        "name": "Server Hardware, Virtualization และ Proxmox VE",
        "competency": "ติดตั้งและบริหารจัดการ Proxmox VE Hypervisor เบื้องต้น",
        "objective": "ข้อ 3 และ 6",
        "content": "Server Hardware; Hypervisor; KVM/LXC; Proxmox; Bridge; Storage; Web UI",
        "activity": "สำรวจ Hardware; ติดตั้ง Proxmox; ตั้ง Network/Storage; ตรวจผ่าน Web UI",
        "media": "Server จริง; Proxmox ISO; USB Boot; Switch; ใบงาน",
        "assessment": "Checklist Proxmox; สาธิตผลจริง; แบบทดสอบ",
        "hours": 5,
    },
    {
        "unit": "5",
        "name": "LXC Container และ Ubuntu Server Stack",
        "competency": "สร้าง Container และติดตั้งแพ็กเก็ตบริการ Web, Database และ Application",
        "objective": "ข้อ 2, 4, 5 และ 6",
        "content": "LXC; Git; Nginx; MariaDB; Node.js; Reverse Proxy; Service Management",
        "activity": "สร้าง Container รายบุคคล; ติดตั้งบริการ; เชื่อม Web/App/DB; ทดสอบจาก Client",
        "media": "Proxmox; Ubuntu Template; Browser; Terminal; คู่มือ Lab",
        "assessment": "ใบงาน Server Stack; ตรวจ Service; ตรวจหน้าเว็บจริง",
        "hours": 5,
    },
    {
        "unit": "6",
        "name": "User, Permission และ SSH Security",
        "competency": "บริหารบัญชีผู้ใช้ สิทธิ์ และการเข้าถึง Server ระยะไกลอย่างปลอดภัย",
        "objective": "ข้อ 2 และ 6",
        "content": "User/Group; chmod; chown; sudo; SSH; Key Authentication; Hardening",
        "activity": "สร้างผู้ใช้ตามสถานการณ์; ทดสอบสิทธิ์; เชื่อม SSH; ตั้งค่า Key และตรวจ Log",
        "media": "Ubuntu Container; SSH Client; Terminal; ใบงาน",
        "assessment": "ภารกิจ Permission; Checklist SSH; สังเกตความปลอดภัย",
        "hours": 5,
    },
    {
        "unit": "7",
        "name": "Server Network, Port และ UFW Firewall",
        "competency": "ตรวจสอบการเชื่อมต่อและกำหนด Firewall Policy ให้บริการ Server",
        "objective": "ข้อ 1, 2 และ 6",
        "content": "IP/Route; Socket/Port; ss; Nmap; UFW; Default Policy; Whitelist; Log; ทบทวน Stack",
        "activity": "ตรวจพอร์ตจากสองฝั่ง; เปิด/ปิด Rule; สแกนด้วย Nmap; วิเคราะห์อาการและ Log",
        "media": "Container รายบุคคล; Nmap; UFW; แผนผัง Port; ใบงาน",
        "assessment": "ใบงาน Firewall; การทดสอบจาก Client; Security Checklist",
        "hours": 10,
    },
    {
        "unit": "8",
        "name": "FTP Server ด้วย vsftpd และ FileZilla",
        "competency": "ติดตั้ง กำหนดสิทธิ์ และทดสอบบริการรับส่งไฟล์ผ่านเครือข่าย",
        "objective": "ข้อ 2, 4 และ 6",
        "content": "FTP; Control/Data Connection; Port 21; Passive Port; vsftpd; User; chroot; Firewall",
        "activity": "ติดตั้ง vsftpd; สร้าง User/Folder; เปิด UFW; Upload/Download; อ่าน Message Log",
        "media": "Proxmox Container; Ubuntu; FileZilla; Animation FTP Flow; คู่มือ Lab",
        "assessment": "ตรวจไฟล์สองฝั่ง; ใบงาน FTP; วิเคราะห์ Troubleshooting",
        "hours": 10,
    },
    {
        "unit": "9",
        "name": "HTTPS, SSL/TLS และ Nginx Reverse Proxy",
        "competency": "เปิดบริการ Web แบบเข้ารหัสและจัดเส้นทางคำขอไปยัง Application",
        "objective": "ข้อ 2, 4, 5 และ 6",
        "content": "TLS/SSL; Certificate; HTTPS Port 443; Reverse Proxy; Nginx Config Test",
        "activity": "ออก Certificate; ตั้ง HTTPS; Proxy ไป Node.js; ทดสอบ Browser และ Log",
        "media": "Nginx; OpenSSL; Node.js App; Browser; ใบงาน",
        "assessment": "ตรวจ HTTPS; ตรวจ Configuration; ใบงาน Reverse Proxy",
        "hours": 5,
    },
    {
        "unit": "10",
        "name": "Docker และการ Deploy Web Application",
        "competency": "Deploy Application ด้วย Container Platform และจัดการทรัพยากรพื้นฐาน",
        "objective": "ข้อ 4, 5 และ 6",
        "content": "Docker Architecture; Image; Container; Port; Volume; Dockerfile; Compose",
        "activity": "ติดตั้ง Docker; Pull/Run Image; สร้าง Dockerfile; Deploy App และทดสอบผ่าน Browser",
        "media": "Docker Engine; Git; Web App; Browser; คู่มือ Lab",
        "assessment": "ตรวจ Container; ตรวจ Web App; ใบงาน Docker",
        "hours": 5,
    },
    {
        "unit": "11",
        "name": "File Sharing, Proxy, AAA และ IoT Platform",
        "competency": "ติดตั้งหรืออธิบายการประยุกต์บริการ Server ตามความต้องการขององค์กร",
        "objective": "ข้อ 4 และ 6",
        "content": "Samba Share; Linux/Samba Permission; Proxy/ACL; AAA; RADIUS; MQTT Broker",
        "activity": "สร้าง Share; ทดสอบจาก Client; วิเคราะห์ ACL; สาธิต Authentication และ MQTT",
        "media": "Samba; Client PC; Proxy Diagram; Mosquitto; MQTT Client",
        "assessment": "ตรวจ Share; ใบงานวิเคราะห์บริการ; แบบทดสอบ",
        "hours": 5,
    },
    {
        "unit": "12",
        "name": "Project บูรณาการระบบเครื่องแม่ข่าย",
        "competency": "วางแผน ติดตั้ง ทดสอบ และจัดทำเอกสารระบบ Server แบบบูรณาการเป็นทีม",
        "objective": "ข้อ 3, 4, 5 และ 6",
        "content": "System Design; Proxmox; Container; Web/DB; Firewall; Service Test; Runbook",
        "activity": "วิเคราะห์โจทย์; แบ่งหน้าที่; สร้างระบบ; ทดสอบ; จัดทำเอกสารและนำเสนอ",
        "media": "Proxmox Lab; Diagram Tool; Git; Template Runbook; Rubric",
        "assessment": "Project Rubric; Demo; Runbook; การทำงานเป็นทีม",
        "hours": 5,
    },
    {
        "unit": "13",
        "name": "ประเมินผลปลายภาคและสรุปสมรรถนะ",
        "competency": "แสดงสมรรถนะการติดตั้งและแก้ปัญหาระบบเครื่องแม่ข่ายตามโจทย์",
        "objective": "ข้อ 1 ถึงข้อ 6",
        "content": "ทบทวนองค์รวม; สอบปลายภาค; สอบปฏิบัติ; Reflection; จิตพิสัย",
        "activity": "สอบข้อเขียน; ปฏิบัติตาม Scenario; อธิบายวิธีตรวจสอบ; สรุปบทเรียน",
        "media": "ข้อสอบ; Proxmox Container; Checklist; แบบประเมินพฤติกรรม",
        "assessment": "ปลายภาค 20; ปฏิบัติ 10; จิตพิสัย 20",
        "hours": 5,
    },
]


def set_run_font(run, size=10, bold=False, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT)
    run.font.complex_script = True
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mm_to_dxa(value):
    return int(round(value * 56.692913))


def set_table_geometry(table, widths_mm):
    widths = [mm_to_dxa(value) for value in widths_mm]
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Mm(widths_mm[index])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.05):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def clear_cell(cell):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    return paragraph


def write_cell(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    paragraph = clear_cell(cell)
    set_paragraph_format(paragraph, align, 0, 0, 1.05)
    run = paragraph.add_run(str(text))
    set_run_font(run, size, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return paragraph


def write_topic_cell(cell, title, topics, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, 0, 1, 1.08)
    r = p.add_run(title)
    set_run_font(r, size + 0.3, True)
    for index, topic in enumerate(topics, start=1):
        p = cell.add_paragraph()
        set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.05)
        r = p.add_run(f"{index}. {topic}")
        set_run_font(r, size, False)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_logo_title(doc, title, subtitle=None, logo_width=0.72):
    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, 0, 1, 1.0)
    p.add_run().add_picture(str(LOGO), width=Inches(logo_width))
    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, 0, 1, 1.0)
    r = p.add_run(title)
    set_run_font(r, 16, True)
    if subtitle:
        p = doc.add_paragraph()
        set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, 0, 3, 1.0)
        r = p.add_run(subtitle)
        set_run_font(r, 10.5, False)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("หน้า ")
    set_run_font(r, 8, False, "666666")
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    r._r.append(fld_char_1)
    r._r.append(instr)
    r._r.append(fld_char_2)


def configure_styles(doc, normal_size=10):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(normal_size)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.05
    for style_name, size in (("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_metadata_table(doc, landscape=False):
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [24, 78, 28, 50] if not landscape else [28, 105, 30, 110]
    set_table_geometry(table, widths)
    metadata = [
        ("ชื่อวิชา", f"{COURSE_NAME} ({COURSE_NAME_EN})", "รหัสวิชา", COURSE_CODE),
        ("จำนวน (ท-ป-น)", COURSE_CREDITS, "ระดับชั้น", LEVEL),
        ("เวลาเรียน", f"{WEEKLY_HOURS} ชั่วโมง : สัปดาห์  รวม {TOTAL_HOURS} ชั่วโมง : ภาคเรียน", "ผู้สอน", TEACHER),
    ]
    for row, values in zip(table.rows, metadata):
        for index, value in enumerate(values):
            label = index in (0, 2)
            write_cell(row.cells[index], value, 9 if landscape else 9.5, label, WD_ALIGN_PARAGRAPH.LEFT)
            if label:
                set_cell_shading(row.cells[index], "E7E6E6")
    return table


def add_course_description_box(doc, landscape=False):
    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, 3, 2, 1.0)
    r = p.add_run("คำอธิบายรายวิชาอ้างอิง")
    set_run_font(r, 10 if landscape else 10.5, True)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [273 if landscape else 180])
    set_cell_shading(table.cell(0, 0), "F2F2F2")
    write_cell(table.cell(0, 0), COURSE_DESCRIPTION, 8.2 if landscape else 9, False, WD_ALIGN_PARAGRAPH.JUSTIFY)


def build_course_project():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(11)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(7)
    configure_styles(doc, 10)
    doc.core_properties.title = f"โครงการสอน {COURSE_CODE} {COURSE_NAME}"
    doc.core_properties.author = TEACHER
    doc.core_properties.subject = SEMESTER

    add_logo_title(
        doc,
        "โครงการสอน/แผนการจัดการเรียนรู้ตลอดภาคเรียน",
        f"{COLLEGE} | {SEMESTER} | 15 สัปดาห์",
        0.72,
    )
    add_metadata_table(doc, landscape=False)
    add_course_description_box(doc, landscape=False)

    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, 4, 2, 1.0)
    r = p.add_run("ตารางกำหนดการสอน 15 สัปดาห์และคะแนนเก็บ")
    set_run_font(r, 11, True)

    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, [14, 14, 118, 18, 16])
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    headers = ["สัปดาห์\nที่", "หน่วย\nที่", "ชื่อหน่วย/รายการสอน", "จำนวน\nชั่วโมง", "คะแนนเก็บ"]
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, 9.3, True, WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    for item in SCHEDULE:
        row = table.add_row()
        set_table_geometry(table, [14, 14, 118, 18, 16])
        write_cell(row.cells[0], item["week"], 9.5, False, WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(row.cells[1], item["unit"], 9.5, False, WD_ALIGN_PARAGRAPH.CENTER)
        write_topic_cell(row.cells[2], item["title"], item["topics"], 8.8)
        write_cell(row.cells[3], item["hours"], 9.5, False, WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(row.cells[4], item["score"], 9.5, False, WD_ALIGN_PARAGRAPH.CENTER)
        set_row_cant_split(row)

    total_row = table.add_row()
    set_table_geometry(table, [14, 14, 118, 18, 16])
    merged = total_row.cells[0].merge(total_row.cells[2])
    write_cell(merged, "รวมจำนวนชั่วโมงและคะแนนตลอดภาคเรียน", 10, True, WD_ALIGN_PARAGRAPH.RIGHT)
    write_cell(total_row.cells[3], sum(item["hours"] for item in SCHEDULE), 10, True, WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(total_row.cells[4], sum(item["score"] for item in SCHEDULE), 10, True, WD_ALIGN_PARAGRAPH.CENTER)
    for cell in total_row.cells:
        set_cell_shading(cell, "E7E6E6")
    set_row_cant_split(total_row)

    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, 5, 2, 1.0)
    r = p.add_run("สรุปสัดส่วนคะแนน")
    set_run_font(r, 11, True)
    score_table = doc.add_table(rows=2, cols=5)
    set_table_geometry(score_table, [36, 36, 36, 36, 36])
    labels = ["คะแนนเก็บ", "สอบปลายภาค", "Project", "จิตพิสัย", "สอบปฏิบัติ"]
    values = [35, 20, 15, 20, 10]
    for index, text in enumerate(labels):
        write_cell(score_table.rows[0].cells[index], text, 9, True, WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(score_table.rows[0].cells[index], "D9E2F3")
        write_cell(score_table.rows[1].cells[index], f"{values[index]} คะแนน", 9.5, True, WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.RIGHT, 8, 0, 1.0)
    r = p.add_run(f"ลงชื่อ ........................................................ ผู้สอน\n({TEACHER})")
    set_run_font(r, 10, False)
    add_page_number(section)

    output = OUT_DOCX / f"โครงการสอน_{COURSE_CODE}_ภาคเรียน1_2569.docx"
    doc.save(output)
    return output


def build_learning_analysis():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(9)
    section.bottom_margin = Mm(10)
    section.left_margin = Mm(12)
    section.right_margin = Mm(12)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(6)
    configure_styles(doc, 8.2)
    doc.core_properties.title = f"ตารางวิเคราะห์การจัดการเรียนรู้ {COURSE_CODE} {COURSE_NAME}"
    doc.core_properties.author = TEACHER
    doc.core_properties.subject = SEMESTER

    add_logo_title(
        doc,
        "ตารางวิเคราะห์การจัดการเรียนรู้",
        f"{COLLEGE} | {SEMESTER} | 15 สัปดาห์",
        0.58,
    )
    add_metadata_table(doc, landscape=True)
    add_course_description_box(doc, landscape=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, 3, 2, 1.0)
    r = p.add_run("การเชื่อมโยงหน่วยการเรียนรู้กับสมรรถนะและการประเมิน")
    set_run_font(r, 10.5, True)

    widths = [12, 32, 38, 37, 45, 44, 24, 28, 13]
    table = doc.add_table(rows=1, cols=9)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    headers = [
        "หน่วย\nที่",
        "ชื่อหน่วย",
        "สมรรถนะวิชาชีพ",
        "จุดประสงค์รายวิชา",
        "เนื้อหาสาระ (โดยย่อ)",
        "กิจกรรมการเรียนรู้",
        "สื่อการเรียนรู้",
        "วิธีการประเมิน",
        "เวลา\n(ชม.)",
    ]
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, 7.8, True, WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")

    for item in ANALYSIS:
        row = table.add_row()
        set_table_geometry(table, widths)
        values = [
            item["unit"],
            item["name"],
            item["competency"],
            item["objective"],
            item["content"],
            item["activity"],
            item["media"],
            item["assessment"],
            item["hours"],
        ]
        for index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if index in (0, 3, 8) else WD_ALIGN_PARAGRAPH.LEFT
            bold = index in (0, 1)
            write_cell(row.cells[index], value, 7.2, bold, align)
        set_row_cant_split(row)

    total_row = table.add_row()
    set_table_geometry(table, widths)
    merged = total_row.cells[0].merge(total_row.cells[7])
    write_cell(merged, "รวมเวลาเรียนตลอดภาคเรียน", 8.3, True, WD_ALIGN_PARAGRAPH.RIGHT)
    write_cell(total_row.cells[8], sum(item["hours"] for item in ANALYSIS), 8.3, True, WD_ALIGN_PARAGRAPH.CENTER)
    for cell in total_row.cells:
        set_cell_shading(cell, "E7E6E6")
    set_row_cant_split(total_row)

    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, 4, 0, 1.0)
    r = p.add_run(
        "หมายเหตุ: สมรรถนะวิชาชีพเชื่อมโยงกับมาตรฐานอาชีพ รหัส 40106 "
        "ช่างสนับสนุนด้านเทคนิค ระดับ 5 และผลลัพธ์การเรียนรู้ระดับรายวิชา"
    )
    set_run_font(r, 8, False, "555555")
    add_page_number(section)

    output = OUT_DOCX / f"ตารางวิเคราะห์การจัดการเรียนรู้_{COURSE_CODE}_ภาคเรียน1_2569.docx"
    doc.save(output)
    return output


if __name__ == "__main__":
    OUT_DOCX.mkdir(parents=True, exist_ok=True)
    project = build_course_project()
    analysis = build_learning_analysis()
    print(project)
    print(analysis)
