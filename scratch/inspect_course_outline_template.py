from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


REFERENCE = Path("/private/tmp/nos-course-outline-template/converted/เค้าโครงการสอน-e1.docx")


def twips_to_mm(value: int | None) -> str:
    if value is None:
        return "auto"
    return f"{value / 56.6929:.2f}"


def row_height_info(row) -> tuple[str, str]:
    tr_pr = row._tr.trPr
    if tr_pr is None:
        return "auto", "auto"
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        return "auto", "auto"
    value = tr_height.get(qn("w:val"))
    rule = tr_height.get(qn("w:hRule")) or "auto"
    return twips_to_mm(int(value)) if value else "auto", rule


def cell_merge_info(cell) -> str:
    tc_pr = cell._tc.tcPr
    parts: list[str] = []
    grid_span = tc_pr.find(qn("w:gridSpan"))
    if grid_span is not None:
        parts.append(f"span={grid_span.get(qn('w:val'))}")
    v_merge = tc_pr.find(qn("w:vMerge"))
    if v_merge is not None:
        parts.append(f"vmerge={v_merge.get(qn('w:val')) or 'continue'}")
    return ",".join(parts) or "-"


def paragraph_properties(paragraph) -> str:
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return ""
    keep_next = p_pr.find(qn("w:keepNext")) is not None
    keep_lines = p_pr.find(qn("w:keepLines")) is not None
    page_break = p_pr.find(qn("w:pageBreakBefore")) is not None
    spacing = p_pr.find(qn("w:spacing"))
    before = spacing.get(qn("w:before")) if spacing is not None else None
    after = spacing.get(qn("w:after")) if spacing is not None else None
    line = spacing.get(qn("w:line")) if spacing is not None else None
    return (
        f"keepNext={keep_next} keepLines={keep_lines} pageBreak={page_break} "
        f"before={before} after={after} line={line}"
    )


def main() -> None:
    doc = Document(REFERENCE)
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")

    print("\nPARAGRAPHS")
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.replace("\t", "<TAB>")
        print(f"P{idx:02d} style={paragraph.style.name!r} {paragraph_properties(paragraph)} text={text!r}")

    for table_idx, table in enumerate(doc.tables):
        print(f"\nTABLE {table_idx}: rows={len(table.rows)} cols={len(table.columns)}")
        grid = table._tbl.tblGrid
        widths = []
        if grid is not None:
            for grid_col in grid.gridCol_lst:
                widths.append(twips_to_mm(int(grid_col.get(qn('w:w')))))
        print(f"grid_mm={widths}")
        for row_idx, row in enumerate(table.rows):
            height, rule = row_height_info(row)
            texts = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
            merges = [cell_merge_info(cell) for cell in row.cells]
            print(
                f"R{row_idx:02d} height_mm={height} rule={rule} "
                f"texts={texts!r} merges={merges!r}"
            )


if __name__ == "__main__":
    main()
