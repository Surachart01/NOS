from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document


REFERENCE = Path("/Users/surachartlimrattanaphun/Downloads/แผนการจัดการเรียนรู้ การใช้เทคโนโลยีดิจิ.docx")


def cell_text(cell) -> str:
    values = [" ".join(paragraph.text.split()) for paragraph in cell.paragraphs if paragraph.text.strip()]
    return " / ".join(values)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--start", type=int, default=None)
    parser.add_argument("--end", type=int, default=None)
    parser.add_argument("--all", action="store_true")
    args = parser.parse_args()
    doc = Document(REFERENCE)
    if args.all or args.start is not None or args.end is not None:
        start = args.start or 0
        end = args.end if args.end is not None else len(doc.paragraphs)
        for index, paragraph in enumerate(doc.paragraphs[start:end], start=start):
            text = " ".join(paragraph.text.split())
            if text:
                print(f"P{index:03d}: {text}")
        return
    print("HEADING-LIKE PARAGRAPHS")
    for index, paragraph in enumerate(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        if (
            re.search(
                r"^(แผนการจัดการเรียนรู้|หน่วยที่|ชื่อหน่วย|หลักสูตรรายวิชา|ตารางวิเคราะห์|โครงการสอน|คำนำ|สารบัญ|ภาคผนวก|ใบงาน|แบบทดสอบ|แบบประเมิน)",
                text,
            )
            or "แผนการจัดการเรียนรู้ที่" in text
            or "บันทึกหลัง" in text
        ):
            print(f"P{index:03d}: {text}")

    print("\nTABLE SUMMARIES")
    for table_index, table in enumerate(doc.tables):
        print(f"\nT{table_index} {len(table.rows)}x{len(table.columns)}")
        for row_index, row in enumerate(table.rows[: min(len(table.rows), 10)]):
            values = [cell_text(cell)[:220] for cell in row.cells]
            print(row_index, values)


if __name__ == "__main__":
    main()
