#!/usr/bin/env python3
"""สร้างเกณฑ์การให้คะแนนสอบปฏิบัติ NOS เป็นไฟล์ Word"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'TH SarabunPSK'
font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')

# Helper: set cell shading
def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color
    })
    shading.append(shading_elm)

# Helper: add styled paragraph to cell
def cell_text(cell, text, bold=False, size=14, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'TH SarabunPSK'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

# ============================================
# TITLE
# ============================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('เกณฑ์การให้คะแนนสอบปฏิบัติ')
run.font.name = 'TH SarabunPSK'
run.font.size = Pt(22)
run.font.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('วิชา ระบบปฏิบัติการบนเครื่องแม่ข่าย (NOS)')
run.font.name = 'TH SarabunPSK'
run.font.size = Pt(18)
run.font.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('วันอังคารที่ 14 กรกฎาคม พ.ศ. 2569  |  คะแนนเต็ม 100 คะแนน')
run.font.name = 'TH SarabunPSK'
run.font.size = Pt(16)
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')

doc.add_paragraph()  # spacer

# ============================================
# ส่วนที่ 1: Nginx (30 คะแนน)
# ============================================
h1 = doc.add_paragraph()
run = h1.add_run('ส่วนที่ 1 — Nginx Web Server (Portfolio)  [30 คะแนน]')
run.font.name = 'TH SarabunPSK'
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(22, 163, 74)
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')

table1 = doc.add_table(rows=1, cols=3)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
table1.style = 'Table Grid'

# Header
for i, text in enumerate(['รายการตรวจ', 'คะแนน', 'ได้']):
    cell = table1.rows[0].cells[i]
    cell_text(cell, text, bold=True, size=15, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, '16A34A')

# Rows
nginx_items = [
    ('ติดตั้ง Nginx สำเร็จ — service ทำงาน', '5'),
    ('มี Heading แสดงชื่อ-นามสกุลของนักเรียน', '5'),
    ('มีข้อมูลส่วนตัว (ชื่อเล่น, อายุ, สาขา, งานอดิเรก ฯลฯ)', '5'),
    ('⭐ มีสโลแกนของตัวเอง (ห้ามขาด!)', '5'),
    ('มีตาราง (Table) อย่างน้อย 1 ตาราง', '3'),
    ('มีลิงก์ (Link) อย่างน้อย 1 ลิงก์ ไปยังเว็บอื่น', '2'),
    ('เว็บเข้าถึงได้ผ่านวง LAN (เครื่องอื่นเปิดดูได้)', '5'),
]

for item, score in nginx_items:
    row = table1.add_row()
    cell_text(row.cells[0], item, size=14)
    cell_text(row.cells[1], score, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[2], '', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

# Total row
row = table1.add_row()
cell_text(row.cells[0], 'รวมส่วนที่ 1', bold=True, size=15, align=WD_ALIGN_PARAGRAPH.RIGHT)
cell_text(row.cells[1], '30', bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_text(row.cells[2], '', size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_shading(row.cells[0], 'DCFCE7')
set_cell_shading(row.cells[1], 'DCFCE7')
set_cell_shading(row.cells[2], 'DCFCE7')

# Set column widths
for row in table1.rows:
    row.cells[0].width = Cm(12)
    row.cells[1].width = Cm(2.5)
    row.cells[2].width = Cm(2.5)

doc.add_paragraph()  # spacer

# ============================================
# ส่วนที่ 2: MariaDB (30 คะแนน)
# ============================================
h2 = doc.add_paragraph()
run = h2.add_run('ส่วนที่ 2 — MariaDB Database  [30 คะแนน]')
run.font.name = 'TH SarabunPSK'
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(37, 99, 235)
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')

table2 = doc.add_table(rows=1, cols=3)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.style = 'Table Grid'

for i, text in enumerate(['รายการตรวจ', 'คะแนน', 'ได้']):
    cell = table2.rows[0].cells[i]
    cell_text(cell, text, bold=True, size=15, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, '2563EB')

mariadb_items = [
    ('ติดตั้ง MariaDB Server สำเร็จ', '10'),
    ('MariaDB Service ทำงานอยู่ (active/running)', '10'),
    ('เข้าหน้า MariaDB ได้ แสดงให้อาจารย์ดู', '10'),
]

for item, score in mariadb_items:
    row = table2.add_row()
    cell_text(row.cells[0], item, size=14)
    cell_text(row.cells[1], score, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[2], '', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

row = table2.add_row()
cell_text(row.cells[0], 'รวมส่วนที่ 2', bold=True, size=15, align=WD_ALIGN_PARAGRAPH.RIGHT)
cell_text(row.cells[1], '30', bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_text(row.cells[2], '', size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_shading(row.cells[0], 'DBEAFE')
set_cell_shading(row.cells[1], 'DBEAFE')
set_cell_shading(row.cells[2], 'DBEAFE')

for row in table2.rows:
    row.cells[0].width = Cm(12)
    row.cells[1].width = Cm(2.5)
    row.cells[2].width = Cm(2.5)

doc.add_paragraph()  # spacer

# ============================================
# ส่วนที่ 3: Firewall (40 คะแนน)
# ============================================
h3 = doc.add_paragraph()
run = h3.add_run('ส่วนที่ 3 — Firewall (UFW)  [40 คะแนน]')
run.font.name = 'TH SarabunPSK'
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(220, 38, 38)
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')

table3 = doc.add_table(rows=1, cols=3)
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
table3.style = 'Table Grid'

for i, text in enumerate(['รายการตรวจ', 'คะแนน', 'ได้']):
    cell = table3.rows[0].cells[i]
    cell_text(cell, text, bold=True, size=15, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, 'DC2626')

# Mission 3.1
row = table3.add_row()
cell_text(row.cells[0], 'Mission 3.1 — เปิดใช้งาน UFW และตั้งค่า Default Policy', bold=True, size=14, color=(220, 38, 38))
cell_text(row.cells[1], '10', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_text(row.cells[2], '', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_shading(row.cells[0], 'FEE2E2')
set_cell_shading(row.cells[1], 'FEE2E2')
set_cell_shading(row.cells[2], 'FEE2E2')

fw_3_1 = [
    ('เปิดใช้งาน UFW สำเร็จ', '4'),
    ('ตั้ง Default deny incoming', '3'),
    ('ตั้ง Default allow outgoing', '3'),
]
for item, score in fw_3_1:
    row = table3.add_row()
    cell_text(row.cells[0], f'    {item}', size=14)
    cell_text(row.cells[1], score, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[2], '', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

# Mission 3.2
row = table3.add_row()
cell_text(row.cells[0], 'Mission 3.2 — จัดการ Port เปิด/ปิด', bold=True, size=14, color=(37, 99, 235))
cell_text(row.cells[1], '10', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_text(row.cells[2], '', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_shading(row.cells[0], 'DBEAFE')
set_cell_shading(row.cells[1], 'DBEAFE')
set_cell_shading(row.cells[2], 'DBEAFE')

fw_3_2 = [
    ('Allow port 22 (SSH)', '2'),
    ('Allow port 80 (HTTP)', '1'),
    ('Allow port 443 (HTTPS)', '1'),
    ('Allow port 53 ทั้ง TCP/UDP (DNS)', '1'),
    ('Deny port 3306 (MariaDB)', '1'),
    ('Deny port 8080', '1'),
    ('Deny port 23 (Telnet)', '1'),
    ('Deny port range 6000:6063/tcp (X11)', '2'),
]
for item, score in fw_3_2:
    row = table3.add_row()
    cell_text(row.cells[0], f'    {item}', size=14)
    cell_text(row.cells[1], score, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[2], '', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

# Mission 3.3
row = table3.add_row()
cell_text(row.cells[0], 'Mission 3.3 — จัดการ IP บล็อค/อนุญาตเฉพาะ', bold=True, size=14, color=(180, 83, 9))
cell_text(row.cells[1], '10', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_text(row.cells[2], '', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_shading(row.cells[0], 'FEF3C7')
set_cell_shading(row.cells[1], 'FEF3C7')
set_cell_shading(row.cells[2], 'FEF3C7')

fw_3_3 = [
    ('บล็อค IP 192.168.1.100', '2'),
    ('บล็อค IP 10.0.0.50', '2'),
    ('อนุญาตเฉพาะ IP 192.168.1.1 เข้า port 22', '3'),
    ('อนุญาต subnet 192.168.1.0/24 เข้า port 80', '3'),
]
for item, score in fw_3_3:
    row = table3.add_row()
    cell_text(row.cells[0], f'    {item}', size=14)
    cell_text(row.cells[1], score, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[2], '', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

# Mission 3.4
row = table3.add_row()
cell_text(row.cells[0], 'Mission 3.4 — ตรวจสอบและ Bonus', bold=True, size=14, color=(22, 163, 74))
cell_text(row.cells[1], '10', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_text(row.cells[2], '', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_shading(row.cells[0], 'DCFCE7')
set_cell_shading(row.cells[1], 'DCFCE7')
set_cell_shading(row.cells[2], 'DCFCE7')

fw_3_4 = [
    ('แสดง ufw status numbered ให้อาจารย์ดู', '3'),
    ('แสดง ufw status verbose เห็น default policy', '3'),
    ('⭐ Bonus: ตั้ง Rate Limiting SSH (limit)', '4'),
]
for item, score in fw_3_4:
    row = table3.add_row()
    cell_text(row.cells[0], f'    {item}', size=14)
    cell_text(row.cells[1], score, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[2], '', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

# Total row
row = table3.add_row()
cell_text(row.cells[0], 'รวมส่วนที่ 3', bold=True, size=15, align=WD_ALIGN_PARAGRAPH.RIGHT)
cell_text(row.cells[1], '40', bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_text(row.cells[2], '', size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_shading(row.cells[0], 'FEE2E2')
set_cell_shading(row.cells[1], 'FEE2E2')
set_cell_shading(row.cells[2], 'FEE2E2')

for row in table3.rows:
    row.cells[0].width = Cm(12)
    row.cells[1].width = Cm(2.5)
    row.cells[2].width = Cm(2.5)

doc.add_paragraph()  # spacer

# ============================================
# GRAND TOTAL
# ============================================
total_table = doc.add_table(rows=1, cols=3)
total_table.alignment = WD_TABLE_ALIGNMENT.CENTER
total_table.style = 'Table Grid'

cell_text(total_table.rows[0].cells[0], 'คะแนนรวมทั้งหมด', bold=True, size=18, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.RIGHT)
cell_text(total_table.rows[0].cells[1], '100', bold=True, size=18, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
cell_text(total_table.rows[0].cells[2], '', bold=True, size=18, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
for cell in total_table.rows[0].cells:
    set_cell_shading(cell, '0F172A')
total_table.rows[0].cells[0].width = Cm(12)
total_table.rows[0].cells[1].width = Cm(2.5)
total_table.rows[0].cells[2].width = Cm(2.5)

doc.add_paragraph()

# ============================================
# ช่องลงนาม
# ============================================
sign = doc.add_paragraph()
sign.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = sign.add_run('ชื่อนักเรียน: ___________________________________    ห้อง: __________    เลขที่: __________')
run.font.name = 'TH SarabunPSK'
run.font.size = Pt(16)
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')

doc.add_paragraph()

sign2 = doc.add_paragraph()
sign2.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = sign2.add_run('หมายเหตุ: _____________________________________________________________________________')
run.font.name = 'TH SarabunPSK'
run.font.size = Pt(14)
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')

sign3 = doc.add_paragraph()
sign3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = sign3.add_run('ลงชื่อผู้ตรวจ ___________________________')
run.font.name = 'TH SarabunPSK'
run.font.size = Pt(14)
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')

# Save
output_path = os.path.expanduser('~/Desktop/NOS/NOS/documents/เกณฑ์การให้คะแนน_สอบปฏิบัติ_NOS.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'✅ สร้างไฟล์เกณฑ์สำเร็จ: {output_path}')
