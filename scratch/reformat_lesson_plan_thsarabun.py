from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(
    "/Users/surachartlimrattanaphun/Desktop/NOS/NOS/output/docx/"
    "แผนการจัดการเรียนรู้_31901-2002_ฉบับละเอียด_ตามแบบฟอร์ม_ปรับหน่วย7_GitHubClone.docx"
)
OUTPUT = Path(
    "/Users/surachartlimrattanaphun/Desktop/NOS/NOS/output/docx/"
    "แผนการจัดการเรียนรู้_31901-2002_ฉบับละเอียด_TH_SarabunPSK_ปรับหน่วย7_GitHubClone.docx"
)

FONT_NAME = "TH SarabunPSK"
BODY_SIZE = Pt(16)
HEADING_SIZE = Pt(18)
COVER_TITLE_SIZE = Pt(24)

TOP_LEVEL_TITLES = {
    "คำนำ",
    "สารบัญ",
    "หลักสูตรรายวิชา",
    "มาตรฐานอาชีพ",
    "ตารางวิเคราะห์หน่วยการเรียนรู้",
    "หน่วยการเรียนรู้",
    "ตารางวิเคราะห์พฤติกรรมการเรียนรู้",
}

LIST_ITEM_RE = re.compile(r"^(?:\d+(?:\.\d+)*|\d+\))\s+")


def set_rfonts(rpr, font_name: str) -> None:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def set_language(rpr) -> None:
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "th-TH")
    lang.set(qn("w:eastAsia"), "th-TH")
    lang.set(qn("w:bidi"), "th-TH")


def set_run_format(run, size) -> None:
    run.font.name = FONT_NAME
    run.font.size = size
    rpr = run._r.get_or_add_rPr()
    set_rfonts(rpr, FONT_NAME)
    set_language(rpr)
    for tag in ("w:spacing", "w:position"):
        node = rpr.find(qn(tag))
        if node is not None:
            rpr.remove(node)
    underline = rpr.find(qn("w:u"))
    if underline is not None and underline.get(qn("w:val")) == "dotted":
        rpr.remove(underline)


def prevent_midword_break(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    word_wrap = ppr.find(qn("w:wordWrap"))
    if word_wrap is None:
        word_wrap = OxmlElement("w:wordWrap")
        ppr.append(word_wrap)
    word_wrap.set(qn("w:val"), "1")


def paragraph_is_bold(paragraph) -> bool:
    visible_runs = [run for run in paragraph.runs if run.text.strip()]
    return bool(visible_runs) and all(bool(run.bold) for run in visible_runs)


def paragraph_has_bold(paragraph) -> bool:
    return any(bool(run.bold) for run in paragraph.runs if run.text.strip())


def is_heading(paragraph, text: str, in_table: bool) -> bool:
    if not text:
        return False
    if text in TOP_LEVEL_TITLES:
        return True
    if not in_table and paragraph_has_bold(paragraph) and len(text) <= 110:
        return True
    if in_table and (
        text == "แผนการจัดการเรียนรู้"
        or re.fullmatch(r"หน่วยที่\s+\d+", text)
        or text.startswith("สอนครั้งที่")
    ):
        return True
    return False


def format_paragraph(
    paragraph,
    *,
    in_table: bool = False,
    body_index: int | None = None,
) -> None:
    for run in paragraph.runs:
        normalized_text = re.sub(r"(?<=\d)-(?=\d)", "‑", run.text)
        normalized_text = normalized_text.replace("สไลสัปดาห์ที่", "สไลด์สัปดาห์ที่")
        normalized_text = normalized_text.replace(
            "; สไลด์สัปดาห์ที่ ",
            ";\nสไลด์สัปดาห์ที่ ",
        )
        if normalized_text != run.text:
            run.text = normalized_text

    text = " ".join(paragraph.text.split())
    heading = is_heading(paragraph, text, in_table)

    if body_index == 1:
        size = COVER_TITLE_SIZE
    elif body_index is not None and 2 <= body_index <= 18 and text:
        size = HEADING_SIZE
    elif heading:
        size = HEADING_SIZE
    else:
        size = BODY_SIZE

    for run in paragraph.runs:
        set_run_format(run, size)

    pf = paragraph.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.widow_control = True
    prevent_midword_break(paragraph)
    if body_index is not None and 77 <= body_index <= 97:
        pf.line_spacing = Pt(19 if heading else 17)
    elif body_index is not None and 41 <= body_index <= 75:
        pf.line_spacing = Pt(17)
    elif in_table:
        pf.line_spacing = Pt(17 if not heading else 19)
    if heading:
        pf.keep_with_next = True
        if (
            not in_table
            and body_index not in range(1, 19)
            and not (body_index is not None and 77 <= body_index <= 97)
        ):
            pf.space_before = Pt(4)
        if paragraph.alignment not in (
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        pf.keep_with_next = False

    if body_index is not None and text in {"สารบัญ", "หลักสูตรรายวิชา"}:
        pf.page_break_before = True

    if not heading and in_table:
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.THAI_JUSTIFY:
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if re.fullmatch(r"[\d./-]+", text)
                else WD_ALIGN_PARAGRAPH.LEFT
            )
    elif not heading and paragraph.alignment == WD_ALIGN_PARAGRAPH.THAI_JUSTIFY:
        if len(text) >= 180 and not LIST_ITEM_RE.match(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif not heading and paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        if len(text) < 180 or LIST_ITEM_RE.match(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def remove_fixed_row_height(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    for height in list(trpr.findall(qn("w:trHeight"))):
        trpr.remove(height)


def set_cell_margins(cell, top=45, start=25, bottom=45, end=25) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.find(qn("w:tcMar"))
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tcmar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_style_font(style, size) -> None:
    style.font.name = FONT_NAME
    style.font.size = size
    rpr = style.element.get_or_add_rPr()
    set_rfonts(rpr, FONT_NAME)
    set_language(rpr)


def set_document_defaults(document) -> None:
    styles = document.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    set_rfonts(rpr, FONT_NAME)
    set_language(rpr)
    for tag, value in (("w:sz", "32"), ("w:szCs", "32")):
        node = rpr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            rpr.append(node)
        node.set(qn("w:val"), value)


def set_table_geometry(table, widths: list[int]) -> None:
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.insert(0, tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    grid_cols = list(grid.findall(qn("w:gridCol")))
    for node in grid_cols:
        grid.remove(node)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for tr in tbl.tr_lst:
        col_index = 0
        for tc in tr.tc_lst:
            tcpr = tc.get_or_add_tcPr()
            grid_span = tcpr.find(qn("w:gridSpan"))
            span = int(grid_span.get(qn("w:val"))) if grid_span is not None else 1
            width = sum(widths[col_index : col_index + span])
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.insert(0, tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            col_index += span


def normalize_table_slashes(table) -> None:
    seen_cells = set()
    for row in table.rows:
        for cell in row.cells:
            if cell._tc in seen_cells:
                continue
            seen_cells.add(cell._tc)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.text = re.sub(r"(?<=[A-Za-z])/(?=[A-Za-z])", " / ", run.text)


def replace_old_outcome_text(document) -> None:
    cell = document.tables[3].rows[10].cells[0]
    if len(cell.paragraphs) < 2:
        return
    cell.paragraphs[1].text = (
        "ผู้เรียนสามารถติดตั้ง กำหนดค่า และตรวจสอบระบบปฏิบัติการเครื่องแม่ข่าย "
        "บริการเครือข่าย ระบบเสมือน และเว็บแอปพลิเคชันบน Ubuntu Server, Proxmox VE "
        "และ LXC ได้อย่างเป็นระบบและปลอดภัย โดยจัดการผู้ใช้ สิทธิ์ SSH Firewall Port "
        "Service Log การสำรองและกู้คืนข้อมูล พร้อมสื่อสารและทำงานร่วมกับผู้อื่นอย่างรับผิดชอบ"
    )


def normalize_empty_page_breaks(document) -> None:
    paragraphs = document.paragraphs
    for break_index, next_index in ((104, 105), (108, 109), (114, 115)):
        paragraph = paragraphs[break_index]
        parent = paragraph._p.getparent()
        if parent is not None:
            parent.remove(paragraph._p)
        paragraphs[next_index].paragraph_format.page_break_before = True


def update_toc(document) -> None:
    toc_pages = {
        41: "ก",
        42: "ข",
        43: "ก",
        44: "ข",
        45: "ค",
        46: "จ",
        47: "ฉ",
        48: "1",
        49: "1",
        50: "4",
        51: "6",
        52: "7",
        53: "7",
        54: "10",
        55: "12",
        56: "13",
        57: "13",
        58: "16",
        59: "19",
        60: "19",
        61: "22",
        62: "24",
        63: "25",
        64: "25",
        65: "28",
        66: "30",
        67: "31",
        68: "31",
        69: "34",
        70: "36",
        72: "37",
        73: "37",
        74: "40",
        75: "42",
    }
    for index, page in toc_pages.items():
        paragraph = document.paragraphs[index]
        label = paragraph.text.rsplit("\t", 1)[0].rstrip()
        paragraph.text = f"{label}\t{page}"


def set_compatibility_settings(document) -> None:
    settings = document.settings.element
    for tag in ("w:autoSpaceDE", "w:autoSpaceDN"):
        node = settings.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            settings.append(node)
        node.set(qn("w:val"), "0")


def iter_header_footer_paragraphs(document):
    seen = set()
    for section in document.sections:
        for part in (section.header, section.first_page_header, section.even_page_header,
                     section.footer, section.first_page_footer, section.even_page_footer):
            key = id(part._element)
            if key in seen:
                continue
            seen.add(key)
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def main() -> None:
    document = Document(SOURCE)
    update_toc(document)
    replace_old_outcome_text(document)
    normalize_empty_page_breaks(document)
    set_document_defaults(document)
    set_compatibility_settings(document)

    for style in document.styles:
        if style.type == 1:
            size = HEADING_SIZE if style.name.startswith("Heading") else BODY_SIZE
            set_style_font(style, size)

    for index, paragraph in enumerate(document.paragraphs):
        format_paragraph(paragraph, body_index=index)

    seen_cells = set()
    for table_index, table in enumerate(document.tables):
        if table_index == 0:
            table.autofit = False
            normalize_table_slashes(table)
            set_table_geometry(table, [850, 1600, 950, 1850, 2044, 2050])
        elif table_index == len(document.tables) - 1:
            table.autofit = False
            set_table_geometry(table, [600, 2400, 3150, 2250])
        else:
            table.autofit = True
        for row in table.rows:
            remove_fixed_row_height(row)
            for cell in row.cells:
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                if table_index == 0:
                    set_cell_margins(cell, top=10, start=25, bottom=10, end=25)
                else:
                    set_cell_margins(cell)
                for paragraph in cell.paragraphs:
                    format_paragraph(paragraph, in_table=True)
                    if table_index == 0:
                        paragraph.paragraph_format.line_spacing = Pt(16)

    for paragraph in iter_header_footer_paragraphs(document):
        format_paragraph(paragraph)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
